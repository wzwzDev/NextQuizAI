#!/usr/bin/env python3
"""
Markdown to DOCX Converter for NextQuizAI Documentation
Converts all .md files to professional Word documents
"""

import os
import sys
from pathlib import Path

# Install required packages if needed
def install_dependencies():
    """Install required Python packages"""
    packages = ['markdown', 'python-docx', 'pypandoc']
    import subprocess
    
    for package in packages:
        try:
            __import__(package.replace('-', '_'))
        except ImportError:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', package, '-q'])

install_dependencies()

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import re

def markdown_to_docx(md_file_path, docx_file_path):
    """Convert Markdown file to DOCX file"""
    
    # Read markdown
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create document
    doc = Document()
    
    # Parse markdown and add to document
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # YAML frontmatter
        if line.strip() == '---':
            # Skip frontmatter
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue
        
        # Headings
        if line.startswith('# '):
            heading = line[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
        
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)
        
        elif line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=4)
        
        # Code blocks
        elif line.strip().startswith('```'):
            # Collect code block
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text, style='List Bullet')
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            i += 1
            continue
        
        # Tables
        elif '|' in line and '-' in lines[i+1] if i+1 < len(lines) else False:
            # Parse table
            table_lines = [line]
            i += 1
            table_lines.append(lines[i])  # separator
            i += 1
            
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            
            # Create table
            rows = len(table_lines) - 2  # exclude header + separator
            cols = len([x for x in table_lines[0].split('|') if x.strip()])
            
            if cols > 0 and rows > 0:
                table = doc.add_table(rows=rows + 1, cols=cols)
                table.style = 'Light Grid Accent 1'
                
                # Header
                header_cells = table.rows[0].cells
                header_items = [x.strip() for x in table_lines[0].split('|') if x.strip()]
                for j, item in enumerate(header_items):
                    header_cells[j].text = item
                
                # Rows
                row_idx = 1
                for line_idx in range(2, len(table_lines)):
                    items = [x.strip() for x in table_lines[line_idx].split('|') if x.strip()]
                    if row_idx < len(table.rows):
                        cells = table.rows[row_idx].cells
                        for j, item in enumerate(items[:cols]):
                            cells[j].text = item
                        row_idx += 1
            
            continue
        
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            doc.add_paragraph(text, style='List Bullet')
        
        # Numbered list
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            doc.add_paragraph(text, style='List Number')
        
        # Blockquotes
        elif line.strip().startswith('> '):
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='Quote')
        
        # Bold & Italic formatting in regular paragraphs
        elif line.strip():
            # Add paragraph with formatting
            text = line.strip()
            p = doc.add_paragraph()
            
            # Simple formatting replacement
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif part:
                    p.add_run(part)
        
        i += 1
    
    # Save document
    doc.save(docx_file_path)
    print(f"✅ Created: {docx_file_path}")

def main():
    """Main conversion function"""
    script_dir = Path(__file__).parent
    
    files_to_convert = [
        '00-objectives.md',
        '01-product-backlog.md',
        '02-sprint1.md',
        '03-inception-deck.md'
    ]
    
    print("🔄 Converting Markdown to DOCX...\n")
    
    for filename in files_to_convert:
        md_path = script_dir / filename
        docx_path = script_dir / filename.replace('.md', '.docx')
        
        if md_path.exists():
            try:
                markdown_to_docx(str(md_path), str(docx_path))
            except Exception as e:
                print(f"❌ Error converting {filename}: {e}")
        else:
            print(f"⚠️  File not found: {filename}")
    
    print("\n✅ ALL CONVERSIONS COMPLETE!")
    print("\n📦 Generated DOCX files:")
    for filename in files_to_convert:
        docx_file = script_dir / filename.replace('.md', '.docx')
        if docx_file.exists():
            size_kb = docx_file.stat().st_size / 1024
            print(f"   ✓ {filename.replace('.md', '.docx')} ({size_kb:.1f} KB)")

if __name__ == '__main__':
    main()
